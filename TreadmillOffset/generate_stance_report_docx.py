import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, color_hex):
    """Set the background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def format_run(run, font_name="Calibri", size_pt=11, bold=False, italic=False, color_rgb=None):
    """Utility to format a text run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_styled_paragraph(doc, text, style_name="Normal", space_after=6, line_spacing=1.15):
    """Add a paragraph with spacing formatting."""
    p = doc.add_paragraph(text, style=style_name)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    return p

def main():
    output_docx = r"d:\AFO_Codes\TreadmillOffset\Treadmill_Stance_Phase_Calculation_Report.docx"
    workspace_docx = r"d:\C3D-parser-1.2.0\Treadmill_Stance_Phase_Calculation_Report.docx"
    
    doc = Document()
    
    # Color palette
    NAVY = RGBColor(27, 54, 93)      # Primary headings
    GREY = RGBColor(100, 100, 100)   # Secondary text
    BLACK = RGBColor(0, 0, 0)
    WHITE = RGBColor(255, 255, 255)
    
    # ----------------------------------------------------
    # TITLE & HEADER
    # ----------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run("GAIT CYCLE STANCE PHASE CALCULATION REPORT")
    format_run(run, size_pt=20, bold=True, color_rgb=NAVY)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run = subtitle_p.add_run("Methodology, Algorithm Description & Subject P03 Results")
    format_run(run, size_pt=12, italic=True, color_rgb=GREY)
    
    # ----------------------------------------------------
    # 1. INTRODUCTION & DEFINITIONS
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    run = h1.add_run("1. Introduction & Definitions")
    format_run(run, size_pt=14, bold=True, color_rgb=NAVY)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "In gait analysis, the gait cycle represents the time interval between two successive occurrences of a repetitive walking event. "
        "Typically, the cycle is defined from the Initial Contact (Heel Strike) of one foot to the subsequent Heel Strike of the same foot. "
        "This gait cycle is divided into two primary phases:"
    )
    
    p = add_styled_paragraph(doc, 
        "• Stance Phase: The period during which the foot remains in contact with the ground. It begins at Initial Contact (Heel Strike) and ends at Foot-Off (Toe-Off).\n"
        "• Swing Phase: The period during which the foot is in the air. It begins at Foot-Off (Toe-Off) and ends at the next Initial Contact (Heel Strike)."
    )
    
    p = add_styled_paragraph(doc, 
        "The Stance Phase Percentage represents the proportion of the entire gait cycle that a limb spends in contact with the ground. "
        "For healthy, symmetric adults walking at customary speeds, the stance phase typically constitutes approximately 60% of the gait cycle, "
        "with the remaining 40% spent in the swing phase. However, this percentage varies depending on factors such as walking speed, incline, and pathology."
    )
    
    # ----------------------------------------------------
    # 2. MATHEMATICAL CALCULATION
    # ----------------------------------------------------
    h2 = doc.add_heading(level=1)
    run = h2.add_run("2. Mathematical Formulation")
    format_run(run, size_pt=14, bold=True, color_rgb=NAVY)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "To calculate the stance phase percentage for a given gait cycle, we identify the exact timestamps of three successive events: "
        "the initial Heel Strike (HS_1), the subsequent Toe-Off (TO), and the next Heel Strike of the same limb (HS_2)."
    )
    
    p = add_styled_paragraph(doc, 
        "The Stance Phase Percentage is computed using the following equation:"
    )
    
    # Equation callout paragraph
    eq_p = doc.add_paragraph()
    eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_p.paragraph_format.space_before = Pt(12)
    eq_p.paragraph_format.space_after = Pt(12)
    run = eq_p.add_run("Stance Phase Percentage (%) = [ (T_TO - T_HS1) / (T_HS2 - T_HS1) ] × 100")
    format_run(run, font_name="Consolas", size_pt=11, bold=True, color_rgb=NAVY)
    
    p = add_styled_paragraph(doc, 
        "Where:\n"
        "• T_HS1 is the time of the initial heel strike (representing 0% of the gait cycle).\n"
        "• T_TO is the time of the subsequent toe-off event.\n"
        "• T_HS2 is the time of the second heel strike (representing 100% of the gait cycle)."
    )
    
    # ----------------------------------------------------
    # 3. ALGORITHM IMPLEMENTATION DETAILS
    # ----------------------------------------------------
    h3 = doc.add_heading(level=1)
    run = h3.add_run("3. Event Detection Algorithm")
    format_run(run, size_pt=14, bold=True, color_rgb=NAVY)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "The automated detection of Heel Strike and Toe-Off events is performed on the Ground Reaction Force (GRF) data. "
        "We utilize the vertical force component (Fy), which corresponds to the column 'ground_force2_vy' for the Right foot and "
        "'ground_force1_vy' for the Left foot. The algorithm is implemented in data_postprocessing.py as follows:"
    )
    
    p = add_styled_paragraph(doc, 
        "1. Heel Strike Detection: The vertical force signal (Fy) is inverted (multiplied by -1). Peak detection is performed "
        "on the inverted signal to find the local minima, which correspond to the swing phases. Starting from each swing valley, "
        "the algorithm searches forward in time to locate the first frame where the vertical force exceeds a threshold of 20 N. "
        "This frame index is designated as a Heel Strike.\n"
        "2. Toe-Off Detection: Peak detection is performed directly on the vertical force (Fy) to find the local maxima, "
        "which correspond to the stance phases. Starting from each stance peak, the algorithm searches forward in time "
        "to locate the first frame where the vertical force drops below the threshold of 20 N. This frame index is designated as a Toe-Off."
    )
    
    # ----------------------------------------------------
    # 4. CRITICAL ROLE OF BASELINE OFFSET CORRECTION
    # ----------------------------------------------------
    h4 = doc.add_heading(level=1)
    run = h4.add_run("4. The Critical Role of Baseline Offset Correction")
    format_run(run, size_pt=14, bold=True, color_rgb=NAVY)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "A critical issue arises in treadmill trials due to baseline sensor drift and belt-rotation noise. "
        "For instance, during the empty-treadmill swing phase, the Left force plate recorded a significant offset of approximately +100 N. "
        "If this baseline offset is not corrected:"
    )
    
    p = add_styled_paragraph(doc, 
        "• The vertical force reading never falls below the 20 N threshold during swing.\n"
        "• The Toe-Off detection search fails, or detects the foot-off event prematurely/belatedly.\n"
        "• This results in highly distorted stance phase calculations (e.g. returning 55% stance with a massive standard deviation of 13.5%)."
    )
    
    p = add_styled_paragraph(doc, 
        "To resolve this, our pipeline applies a Treadmill Offset Corrector to tare out belt-rotation offsets and subsequently runs a "
        "dynamic baseline filter (baseline_correct_debug). This filter identifies swing valleys, interpolates the offset curve, and "
        "subtracts it from the signal. This tares the force plate output to exactly 0 N during the swing phase, ensuring that the "
        "20 N threshold is crossed cleanly and yielding highly accurate event detection."
    )
    
    # ----------------------------------------------------
    # 5. SUBJECT P03 GAIT ANALYSIS RESULTS
    # ----------------------------------------------------
    h5 = doc.add_heading(level=1)
    run = h5.add_run("5. Subject P03 Gait Analysis Results")
    format_run(run, size_pt=14, bold=True, color_rgb=NAVY)
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "Using the baseline-corrected data, we recalculated the stance phase percentages for all trials under the K3 dataset. "
        "The average stance percentages for the Right and Left feet are summarized in the table below:"
    )
    
    # Add Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Trial File Name"
    hdr_cells[1].text = "Right Foot Stance (%)"
    hdr_cells[2].text = "Left Foot Stance (%)"
    
    for cell in hdr_cells:
        set_cell_background(cell, "1B365D")
        for p_cell in cell.paragraphs:
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r_cell in p_cell.runs:
                format_run(r_cell, size_pt=10, bold=True, color_rgb=WHITE)
                
    results_data = [
        ("k3 speed test.mot", "69.87%", "65.21%"),
        ("k3 slope 1.mot", "72.86%", "71.48%"),
        ("k3 slope actual 2.mot", "72.59%", "69.41%"),
        ("k3 slope actual 3.mot", "72.55%", "67.03%"),
    ]
    
    for trial_name, right_val, left_val in results_data:
        row_cells = table.add_row().cells
        row_cells[0].text = trial_name
        row_cells[1].text = right_val
        row_cells[2].text = left_val
        
        # Center align and style values
        for cell in row_cells[1:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    p = add_styled_paragraph(doc, "") # Spacing after table
    
    p = add_styled_paragraph(doc, 
        "Key Clinical Findings & Observations:\n"
        "• Prolonged Stance Phase: For both feet, the stance phase percentage is systematically higher than the standard 60% mark. "
        "This indicates a strategy of keeping the feet on the ground longer (increased double support time), which is commonly observed "
        "in individuals seeking increased stability or adapting to treadmill walking.\n"
        "• Asymmetric Gait: The Right foot exhibits a systematically longer stance phase (69.9% – 72.9%) than the Left foot (65.2% – 71.5%). "
        "This asymmetry is clinically relevant and typically indicates that the participant is loading or preferring the right side "
        "or that the left side has a faster swing-through or shorter stance time (which might occur if walking with an orthosis like an AFO, "
        "or due to weakness/pain on one side).\n"
        "• Ankle Power Peak Alignment: In the biomechanics plots, the concentric ankle power generation peak (representing the push-off energy) "
        "occurs at approximately 55% - 60% of the gait cycle. Correctly placing the toe-off boundaries (at 65.2% on the left and 69.9% on the right) "
        "successfully places this ankle power peak inside the stance phase (prior to the toe-off line), confirming physiological validity."
    )
    
    doc.save(output_docx)
    doc.save(workspace_docx)
    print("Word report generated and saved at:", output_docx)
    print("Word report also copied to workspace at:", workspace_docx)

if __name__ == "__main__":
    main()
