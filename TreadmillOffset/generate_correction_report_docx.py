import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, color_hex):
    """Set the background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def format_run(run, font_name="Calibri", size_pt=11, bold=False, italic=False, color_rgb=None):
    """Utility to format a text run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_styled_paragraph(doc, text, style_name="Normal", space_after=6, line_spacing=1.15):
    """Add a paragraph with spacing formatting."""
    p = doc.add_paragraph(text, style=style_name)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    return p

def main():
    output_docx = r"d:\AFO_Codes\TreadmillOffset\Treadmill_Offset_Correction_Methodology.docx"
    workspace_docx = r"d:\C3D-parser-1.2.0\Treadmill_Offset_Correction_Methodology.docx"
    
    doc = Document()
    
    # Color palette
    NAVY = RGBColor(27, 54, 93)      # Primary headings
    GREY = RGBColor(100, 100, 100)   # Secondary text
    BLACK = RGBColor(0, 0, 0)
    
    # ----------------------------------------------------
    # TITLE & HEADER
    # ----------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run("TREADMILL FORCE PLATE OFFSET CORRECTION")
    format_run(run, size_pt=24, bold=True, color_rgb=NAVY)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run = subtitle_p.add_run("Methodology & Software Implementation Details")
    format_run(run, size_pt=14, italic=True, color_rgb=GREY)
    
    # ----------------------------------------------------
    # 1. OVERVIEW & RATIONALE
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    run = h1.add_run("1. Overview & Rationale")
    format_run(run, size_pt=16, bold=True, color_rgb=NAVY)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "During gait analysis on instrumented treadmills, force plates record contact forces (ground reaction forces) from walking. "
        "However, even when no one is walking on the belt (empty-belt condition), the rotation of the treadmill belt, mechanical friction, "
        "and vibration cause non-zero baseline forces to be recorded. These baseline force plate offsets vary non-linearly with "
        "treadmill speed and incline slope."
    )
    
    p = add_styled_paragraph(doc, 
        "To prevent systematic errors in joint moment and joint power calculations during inverse dynamics, these empty-belt offsets "
        "must be tared out. Since the offset values change dynamically with speed and slope, a simple static subtraction is insufficient. "
        "Therefore, we implemented a custom, automated calibration and correction software system."
    )
    
    # ----------------------------------------------------
    # 2. SOFTWARE ARCHITECTURE
    # ----------------------------------------------------
    h2 = doc.add_heading(level=1)
    run = h2.add_run("2. Software Architecture & The Corrector Class")
    format_run(run, size_pt=16, bold=True, color_rgb=NAVY)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "The core correction logic is implemented in the class TreadmillOffsetCorrector (located in d:\\AFO_Codes\\TreadMetrix\\offset_corrector.py). "
        "This class acts as both a Lookup Table (for exact speed/slope matches) and a 2D Grid Interpolator for conditions falling between collected data points."
    )
    
    p = add_styled_paragraph(doc, 
        "The correction process consists of three main stages:\n"
        "1. Initialization: The corrector reads the pooled calibration summary CSV (pooled_treadmill_offsets.csv) by default. This pooled dataset combines "
        "and averages the offsets from both Day01 and Day02 for overlapping incline conditions (such as the 3.1% slope trials) and preserves unique incline trials "
        "(e.g., slopes of 0.0% to 4.5% from Day01 and slopes of 5.0% and 7.0% from Day02), resulting in a single unified 120-trial calibration sweep.\n"
        "2. 2D Surface Fitting: For each of the 6 force channels (3 axes on Plate 4, 3 axes on Plate 5), the corrector fits a continuous "
        "triangulated surface using scipy.interpolate.LinearNDInterpolator over the coordinates of (Speed, Slope).\n"
        "3. Offset Estimation: For any input query of (Speed, Slope), the corrector evaluates the fitted 2D surface to obtain the taring offsets. "
        "If a query falls outside the calibration limits (extrapolation), it automatically performs a nearest-neighbor lookup to prevent NaN errors."
    )
    
    # ----------------------------------------------------
    # 3. IMPLEMENTED CORRECTION METHODS
    # ----------------------------------------------------
    h3 = doc.add_heading(level=1)
    run = h3.add_run("3. Implemented Correction Methods")
    format_run(run, size_pt=16, bold=True, color_rgb=NAVY)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "Depending on the experimental protocol, the corrector supports two distinct taring methods:"
    )
    
    p = add_styled_paragraph(doc, 
        "Method A: Whole-Trial Global Correction (correct_mot_dataframe)\n"
        "This method applies if the entire trial was recorded at a single, constant speed and incline slope. "
        "The function loads the trial's .mot file into a Pandas DataFrame, queries the 2D surface interpolator for the corresponding offsets, "
        "and subtracts the offsets globally from all data rows. It also handles mapping between calibration force plate labels (Plates 4 and 5) "
        "and the experimental trial labels (Plates 1 and 2, which correspond to Left and Right feet)."
    )
    
    p = add_styled_paragraph(doc, 
        "Method B: Interactive Piecewise Window Correction (interactive_correction)\n"
        "For multi-speed trials recorded continuously in a single file, the offset changes at different times. "
        "This method opens a graphical user interface (GUI) built with Matplotlib and Tkinter:\n"
        "• The user is shown a plot of the trial's vertical force over time and drags a span selector to highlight steady-speed windows.\n"
        "• A pop-up dialog prompts the user to enter the specific speed (mph) and slope (%) for that highlighted window.\n"
        "• The script queries the 2D surface models and subtracts the corresponding offsets only within that highlighted time window.\n"
        "• The workflow returns both the corrected data and the list of window parameters so that subsequent steps can route "
        "segmented gait cycles into the correct speed-labeled subfolders."
    )
    
    # ----------------------------------------------------
    # 4. DATA PROCESSING PIPELINE INTEGRATION
    # ----------------------------------------------------
    h4 = doc.add_heading(level=1)
    run = h4.add_run("4. Data Processing Pipeline Integration")
    format_run(run, size_pt=16, bold=True, color_rgb=NAVY)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "The corrector is integrated directly into the main pipeline script (full_pipeline.py):\n"
        "1. Load Files: Raw .mot and .trc files are loaded into a Trial object.\n"
        "2. Offset Tare: The interactive_correction method is invoked to apply speed-specific offsets.\n"
        "3. Post-Processing: data_postprocessing.py filters the data, performs dynamic swing-phase baseline taring to remove minor sensor drifts, "
        "detects heel strikes and toe-offs, and segments the trial into individual gait cycles.\n"
        "4. Output Routing: Based on the selected time windows, gait cycles are sorted by speed/slope (e.g. 'Speed1_25slope3_1') and routed "
        "into respective folders for Inverse Kinematics (IK), Inverse Dynamics (ID), and Joint Power (JP) calculations."
    )
    
    # ----------------------------------------------------
    # 5. CODE SNIPPET REFERENCE
    # ----------------------------------------------------
    h5 = doc.add_heading(level=1)
    run = h5.add_run("5. Code Implementation Summary")
    format_run(run, size_pt=16, bold=True, color_rgb=NAVY)
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "Below is a conceptual layout of the correction application in the post-processing script:"
    )
    
    # Add table for conceptual layout
    table_code = doc.add_table(rows=1, cols=2)
    table_code.style = 'Light Shading Accent 1'
    hdr_cells = table_code.rows[0].cells
    hdr_cells[0].text = "Correction Component"
    hdr_cells[1].text = "Description & Implementation Details"
    
    for cell in hdr_cells:
        set_cell_background(cell, "1B365D")
        for p_cell in cell.paragraphs:
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r_cell in p_cell.runs:
                format_run(r_cell, size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
                
    components = [
        ("TreadmillOffsetCorrector()", "Loads the pooled calibration CSV (pooled_treadmill_offsets.csv) and initializes scipy's LinearNDInterpolator for Plates 4 & 5."),
        ("get_offsets(speed, slope)", "Computes the 6 force plate offset values. Falls back to nearest neighbor if out-of-bounds."),
        ("correct_mot_dataframe(df, speed, slope)", "Subtracts the 6 offsets globally from the target columns in the Pandas DataFrame."),
        ("interactive_correction(df)", "Opens Matplotlib/Tkinter GUI to apply piecewise offsets only within selected time bounds.")
    ]
    
    for comp, desc in components:
        row_cells = table_code.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = desc
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        
    doc.save(output_docx)
    doc.save(workspace_docx)
    print("Word report generated and saved at:", output_docx)
    print("Word report also copied to workspace at:", workspace_docx)

if __name__ == "__main__":
    main()
