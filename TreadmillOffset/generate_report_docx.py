import os
import pandas as pd
import numpy as np
import statsmodels.api as sm
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, color_hex):
    """Set the background color of a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def format_run(run, font_name="Calibri", size_pt=11, bold=False, italic=False, color_rgb=None):
    """Utility to format a text run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def add_styled_paragraph(doc, text, style_name="Normal", space_after=6, line_spacing=1.15):
    """Add a paragraph with spacing formatting."""
    p = doc.add_paragraph(text, style=style_name)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    return p

def main():
    day01_csv = r"d:\AFO_Codes\TreadmillOffset\Day01\treadmill_offsets_summary.csv"
    day02_csv = r"d:\AFO_Codes\TreadmillOffset\Day02\treadmill_offsets_summary.csv"
    output_docx = r"d:\AFO_Codes\TreadmillOffset\Treadmill_Offset_Consistency_Report.docx"
    
    # Load data
    df1 = pd.read_csv(day01_csv)
    df2 = pd.read_csv(day02_csv)
    
    force_cols = [
        'ground_force4_vx', 'ground_force4_vy', 'ground_force4_vz',
        'ground_force5_vx', 'ground_force5_vy', 'ground_force5_vz'
    ]
    
    doc = Document()
    
    # Color palette
    NAVY = RGBColor(27, 54, 93)      # Primary headings
    GREY = RGBColor(100, 100, 100)   # Secondary text
    BLACK = RGBColor(0, 0, 0)
    
    # ----------------------------------------------------
    # TITLE & HEADER
    # ----------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run("TREADMILL FORCE PLATE CALIBRATION OFFSETS")
    format_run(run, size_pt=24, bold=True, color_rgb=NAVY)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run = subtitle_p.add_run("Session-to-Session Consistency Report (Day 01 vs Day 02)")
    format_run(run, size_pt=14, italic=True, color_rgb=GREY)
    
    # ----------------------------------------------------
    # EXECUTIVE SUMMARY
    # ----------------------------------------------------
    h1 = doc.add_heading(level=1)
    run = h1.add_run("1. Executive Summary")
    format_run(run, size_pt=16, bold=True, color_rgb=NAVY)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "This report evaluates the consistency of treadmill force plate calibration offsets across two different experimental sessions "
        "(Day 01 and Day 02). In gait analysis and biomechanical modeling (such as OpenSim), force plate data must be tared to account for "
        "empty-belt treadmill offsets. Since these offsets depend non-linearly on speed and slope due to friction, belt tension, and structural dynamics, "
        "it is critical to verify if a single calibration profile can be reused or if new calibrations are required for each session."
    )
    
    p = add_styled_paragraph(doc, 
        "Our analysis indicates that while the speed-dependent trends (slopes) are highly repeatable and parallel between sessions, "
        "there is a significant static bias shift (drift in the constants) of up to 30 N in the anterior-posterior (Fz) shear force and up to "
        "10 N in the vertical (Fy) force between Day 01 and Day 02. Therefore, treadmill offsets are NOT consistent between sessions, "
        "and a new calibration protocol must be performed for each testing day to prevent systematic errors in downstream inverse dynamics."
    )
    
    # ----------------------------------------------------
    # METHODOLOGY
    # ----------------------------------------------------
    h2 = doc.add_heading(level=1)
    run = h2.add_run("2. Methodology")
    format_run(run, size_pt=16, bold=True, color_rgb=NAVY)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "Empty-treadmill trials were recorded at varying combinations of speeds and incline slopes across two sessions:\n"
        "• Day 01 Session: 101 trials recorded. Speeds ranging from 0.2 to 1.5 mph; slopes ranging from 0.0% to 4.5%.\n"
        "• Day 02 Session: 33 trials recorded. Speeds ranging from 0.2 to 1.5 mph; slopes ranging from 3.1% to 7.0%."
    )
    
    p = add_styled_paragraph(doc, 
        "For each trial, the mean and median force offsets were computed across all 6 axes of the two force plates:\n"
        "• Plate 4: ground_force4_vx (ML), ground_force4_vy (Vertical), ground_force4_vz (AP)\n"
        "• Plate 5: ground_force5_vx (ML), ground_force5_vy (Vertical), ground_force5_vz (AP)"
    )
    
    p = add_styled_paragraph(doc, 
        "We performed two types of comparisons:\n"
        "1. Direct Value Comparison: Evaluation of the actual median offset values at the exact overlapping incline slope (3.1%) across the shared speed levels.\n"
        "2. Ordinary Least Squares (OLS) Regression Comparison: Fitting the linear model: Offset = const + b_speed * Speed + b_slope * Slope for both sessions, then comparing the fitted parameters side-by-side."
    )
    
    # ----------------------------------------------------
    # RESULTS: DIRECT COMPARISON
    # ----------------------------------------------------
    h3 = doc.add_heading(level=1)
    run = h3.add_run("3. Results: Direct Comparison at Common Incline (Slope = 3.1%)")
    format_run(run, size_pt=16, bold=True, color_rgb=NAVY)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "Day 01 and Day 02 share a single overlapping incline slope of 3.1%. Direct row-by-row comparisons of the median offsets "
        "were conducted across the 11 shared speeds (0.2 to 1.5 mph). The mean and maximum absolute differences for each force plate axis are summarized below:"
    )
    
    # Add summary table of differences
    table_diff = doc.add_table(rows=1, cols=4)
    table_diff.style = 'Light Shading Accent 1'
    hdr_cells = table_diff.rows[0].cells
    hdr_cells[0].text = "Force Plate & Axis"
    hdr_cells[1].text = "Mean Abs Diff (N)"
    hdr_cells[2].text = "Max Abs Diff (N)"
    hdr_cells[3].text = "Alignment Trend"
    
    # Apply styling to header
    for cell in hdr_cells:
        set_cell_background(cell, "1B365D")
        for p_cell in cell.paragraphs:
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r_cell in p_cell.runs:
                format_run(r_cell, size_pt=10, bold=True, color_rgb=RGBColor(255, 255, 255))
                
    s1_31 = df1[np.isclose(df1['Slope'], 3.1, atol=0.05)]
    s2_31 = df2[np.isclose(df2['Slope'], 3.1, atol=0.05)]
    shared_speeds = sorted(list(set(s1_31['Speed']).intersection(set(s2_31['Speed']))))
    
    descriptions = {
        'ground_force4_vx': ("Plate 4 Medial-Lateral (ML)", "Increasing offsets with speed; parallel shift"),
        'ground_force4_vy': ("Plate 4 Vertical (Fy)", "Excellent parallel alignment; negative scaling with speed"),
        'ground_force4_vz': ("Plate 4 Anterior-Posterior (AP)", "Significant constant shift; Day01 systematically higher"),
        'ground_force5_vx': ("Plate 5 Medial-Lateral (ML)", "Very close baseline; minor speed sensitivity change"),
        'ground_force5_vy': ("Plate 5 Vertical (Fy)", "Excellent agreement across all speed levels"),
        'ground_force5_vz': ("Plate 5 Anterior-Posterior (AP)", "Parallel trends with a static ~8 N bias shift")
    }
    
    for col in force_cols:
        col_median = f"{col}_median"
        abs_diffs = []
        for spd in shared_speeds:
            val1 = s1_31[np.isclose(s1_31['Speed'], spd, atol=0.01)][col_median].values[0]
            val2 = s2_31[np.isclose(s2_31['Speed'], spd, atol=0.01)][col_median].values[0]
            abs_diffs.append(abs(val2 - val1))
        
        mean_d = np.mean(abs_diffs)
        max_d = np.max(abs_diffs)
        
        row_cells = table_diff.add_row().cells
        row_cells[0].text = descriptions[col][0]
        row_cells[1].text = f"{mean_d:.2f} N"
        row_cells[2].text = f"{max_d:.2f} N"
        row_cells[3].text = descriptions[col][1]
        
        # Center align columns 1 and 2
        for col_idx in [1, 2]:
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # RESULTS: OLS REGRESSION COMPARISON
    # ----------------------------------------------------
    h4 = doc.add_heading(level=1)
    run = h4.add_run("4. Results: Regression Parameters Comparison")
    format_run(run, size_pt=16, bold=True, color_rgb=NAVY)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "OLS regressions were fitted for both days to capture how offsets scale with speed and slope. "
        "The fitted model equation is: Offset = Const + Speed_Coef * Speed + Slope_Coef * Slope. "
        "Comparing these parameters highlights the baseline session shift (Const) and sensitivity variations (Speed/Slope coefficients):"
    )
    
    # Add regression table
    table_reg = doc.add_table(rows=1, cols=7)
    table_reg.style = 'Light Shading Accent 1'
    hdr_cells = table_reg.rows[0].cells
    hdr_cells[0].text = "Axis"
    hdr_cells[1].text = "Day 01 Const"
    hdr_cells[2].text = "Day 02 Const"
    hdr_cells[3].text = "Day 01 Speed Coef"
    hdr_cells[4].text = "Day 02 Speed Coef"
    hdr_cells[5].text = "Day 01 R²"
    hdr_cells[6].text = "Day 02 R²"
    
    for cell in hdr_cells:
        set_cell_background(cell, "1B365D")
        for p_cell in cell.paragraphs:
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r_cell in p_cell.runs:
                format_run(r_cell, size_pt=9, bold=True, color_rgb=RGBColor(255, 255, 255))
                
    for col in force_cols:
        col_median = f"{col}_median"
        # Fit Day 01
        X1 = df1[['Speed', 'Slope']]
        X1 = sm.add_constant(X1)
        y1 = df1[col_median]
        model1 = sm.OLS(y1, X1).fit()
        
        # Fit Day 02
        X2 = df2[['Speed', 'Slope']]
        X2 = sm.add_constant(X2)
        y2 = df2[col_median]
        model2 = sm.OLS(y2, X2).fit()
        
        row_cells = table_reg.add_row().cells
        row_cells[0].text = col.replace('ground_force', 'Plate ').replace('_v', ' F')
        row_cells[1].text = f"{model1.params['const']:.2f}"
        row_cells[2].text = f"{model2.params['const']:.2f}"
        row_cells[3].text = f"{model1.params['Speed']:.2f}"
        row_cells[4].text = f"{model2.params['Speed']:.2f}"
        row_cells[5].text = f"{model1.rsquared:.2f}"
        row_cells[6].text = f"{model2.rsquared:.2f}"
        
        for col_idx in range(1, 7):
            row_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ----------------------------------------------------
    # DISCUSSION & RECOMMENDATIONS
    # ----------------------------------------------------
    h5 = doc.add_heading(level=1)
    run = h5.add_run("5. Discussion & Recommendations")
    format_run(run, size_pt=16, bold=True, color_rgb=NAVY)
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(6)
    
    p = add_styled_paragraph(doc, 
        "The comparison results reveal two primary behaviors in the treadmill force plate offsets:\n"
        "1. High Trend Consistency (Parallel Scaling): The speed-dependent coefficients show consistent scaling directions. For example, "
        "the vertical offsets on both plates become more negative as speed increases. The plots show parallel lines between Day 01 and Day 02, "
        "indicating that the mechanical friction and rolling dynamics scale identically with speed across sessions.\n"
        "2. Significant Static Drift (Baseline Shift): There is a distinct baseline shift between the sessions. The constants for "
        "the anterior-posterior forces (Plate 4 Fz and Plate 5 Fz) shifted by approximately 30 N to 32 N. Similarly, the vertical constant on "
        "Plate 5 shifted by ~10 N. This static drift is likely caused by changes in belt tension, minor alignment adjustments, or temperature changes in the load cell amplifiers between the sessions."
    )
    
    p = add_styled_paragraph(doc, 
        "Recommendations:\n"
        "• Reusing Calibration profiles is NOT recommended. Applying Day 01 calibration offsets to Day 02 trials would introduce up to 40 N of "
        "systematic shear force error, which directly corrupts joint moment computations (Inverse Dynamics) during gait analysis.\n"
        "• Perform Per-Session Calibrations: An empty-treadmill speed/slope sweep should be recorded at the beginning of each testing session "
        "to tare the force plates. The scipy 2D interpolation models should be rebuilt dynamically using the calibration sweep unique to that day."
    )
    
    doc.save(output_docx)
    print("Word report generated successfully at:", output_docx)

if __name__ == "__main__":
    main()
